import os
import re
import json
import time
import random
from django.shortcuts import render, redirect
from django.http import JsonResponse
from django.contrib.auth import login, authenticate
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm, PasswordResetForm
from django.contrib import messages
from django.core.cache import cache
from django.db import transaction
from .models import UserInfo
from docx import Document
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth import get_user_model
from django.contrib.auth.hashers import make_password

#-----------------------------通用模块 登录注册注销退出-----------------------------#

def home(request):
    """系统首页视图"""
    login_form = AuthenticationForm()
    return render(request, 'registration/index.html', {
        'login_form': login_form,
    })

def login_view(request):
    """用户登录视图(支持手机验证码和手机号+密码登录)"""
    if request.method == 'POST':
        # 手机验证码登录
        if 'phone' in request.POST and 'verification_code' in request.POST:
            phone = request.POST.get('phone')
            code = request.POST.get('verification_code')
            cached_code = cache.get(f'security_code_{phone}_login')
            
            if not cached_code:
                messages.error(request, '请先获取验证码')
                return redirect('login')
                
            if cached_code != code:
                messages.error(request, '验证码错误')
                return render(request, 'registration/index.html', {
                    'login_form': AuthenticationForm(),
                    'phone': phone,
                    'verification_error': True
                })
            
            User = get_user_model()
            try:
                user = User.objects.get(phone=phone)
                if not user.has_usable_password():
                    pass
            except User.DoesNotExist:
                user = User.objects.create_user(phone=phone)
                avatar_num = random.randint(1, 6)
                UserInfo.objects.create(
                    user=user,
                    nickname=f'用户{phone[-4:]}',
                    gender='未知',
                    avatar=f'avatars/default_{avatar_num}.png'
                )
                
            login(request, user)
            if hasattr(user, 'role'):
                if user.role == 'B':
                    return redirect('provider_dashboard')
                return redirect('user_dashboard')
            return redirect('home')
        
        # 手机号+密码登录
        elif 'phone' in request.POST and 'password' in request.POST:
            phone = request.POST.get('phone')
            password = request.POST.get('password')
            
            if not re.match(r'^1[3-9]\d{9}$', phone):
                messages.error(request, '手机号格式不正确')
                return redirect('login')
                
            User = get_user_model()
            try:
                user = User.objects.get(phone=phone)
                if not user.has_usable_password():
                    messages.error(request, '该账号未设置密码，请使用验证码登录')
                    return redirect('login')
                    
                if not user.check_password(password):
                    messages.error(request, '手机号或密码错误')
                    return redirect('login')
                    
                login(request, user)
                if hasattr(user, 'role'):
                    if user.role == 'B':
                        return redirect('provider_dashboard')
                    return redirect('user_dashboard')
                return redirect('home')
                
            except User.DoesNotExist:
                messages.error(request, '手机号未注册')
                return redirect('login')
        
        # 传统用户名密码登录
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)
            if user is not None:
                login(request, user)
                if hasattr(user, 'role'):
                    if user.role == 'provider':
                        return redirect('provider_dashboard')
                    return redirect('user_dashboard')
                return redirect('home')
        messages.error(request, '用户名或密码错误')
    return redirect('home')

def terms_view(request):
    """用户协议视图"""
    doc_path = os.path.join(settings.BASE_DIR, 'core/static/docs/terms.docx')
    doc = Document(doc_path)
    content = '\n'.join([para.text for para in doc.paragraphs])
    return render(request, 'registration/terms.html', {'content': content})

def privacy_view(request):
    """隐私政策视图""" 
    doc_path = os.path.join(settings.BASE_DIR, 'core/static/docs/privacy.docx')
    doc = Document(doc_path)
    content = '\n'.join([para.text for para in doc.paragraphs])
    return render(request, 'registration/privacy.html', {'content': content})

@csrf_exempt
def send_verification_code(request):
    """发送验证码API(开发模式)"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            phone = data.get('phone')
            purpose = data.get('purpose', 'login')  # login/reset/change_phone
            
            if not phone or not re.match(r'^1[3-9]\d{9}$', phone):
                return JsonResponse({'status': 'error', 'message': '手机号无效'}, status=400)
            
            # 生成6位随机验证码
            code = ''.join([str(random.randint(0, 9)) for _ in range(6)])
            # 存储验证码到缓存，有效期5分钟
            cache.set(f'security_code_{phone}_{purpose}', code, 300)
            
            return JsonResponse({
                'status': 'success',
                'message': '验证码已发送(开发模式)',
                'code': code,  # 开发模式下返回验证码
                'countdown': 60
            })
            
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    
    return JsonResponse({'status': 'error', 'message': '仅支持POST请求'}, status=405)

#-----------------------------用户模块-----------------------------#

def user_dashboard(request):
    """普通用户仪表盘视图"""
    if not request.user.is_authenticated:
        return redirect('login')
    try:
        user_info = UserInfo.objects.get(user=request.user)
    except UserInfo.DoesNotExist:
        avatar_num = random.randint(1, 6)
        user_info = UserInfo.objects.create(
            user=request.user,
            nickname=f'用户{request.user.username[:4]}',
            gender='未知',
            avatar=f'avatars/default_{avatar_num}.png'
        )
    return render(request, 'user/common/dashboard.html', {
        'user': request.user,
        'user_info': user_info
    })

def provider_dashboard(request):
    """服务提供者仪表盘视图"""
    if not request.user.is_authenticated:
        return redirect('login')
    if not hasattr(request.user, 'role') or request.user.role != 'provider':
        return redirect('user_dashboard')
    return render(request, 'provider\common\dashboard.html', {
        'user': request.user
    })

#-----------------------------安全设置API-----------------------------#

def change_phone_api(request):
    """更换手机号API"""
    try:
        data = json.loads(request.body)
        old_phone = request.user.phone
        new_phone = data.get('new_phone')
        code = data.get('code')
        
        if not new_phone or not re.match(r'^1[3-9]\d{9}$', new_phone):
            return JsonResponse({'status': 'error', 'message': '请输入有效的新手机号码'}, status=400)
            
        if not code or not re.match(r'^\d{6}$', code):
            return JsonResponse({'status': 'error', 'message': '请输入6位验证码'}, status=400)
            
        # 验证原手机号验证码
        cached_code = cache.get(f'security_code_{old_phone}_change_phone')
        if not cached_code or cached_code != code:
            return JsonResponse({'status': 'error', 'message': '验证码错误或已过期'}, status=400)
            
        # 检查新手机号是否已被使用
        User = get_user_model()
        if User.objects.filter(phone=new_phone).exclude(id=request.user.id).exists():
            return JsonResponse({'status': 'error', 'message': '该手机号已被其他账号使用'}, status=400)
            
        # 更新手机号
        request.user.phone = new_phone
        request.user.save()
        
        # 清除验证码缓存
        cache.delete(f'security_code_{old_phone}_change_phone')
        
        return JsonResponse({
            'status': 'success', 
            'message': '手机号更换成功',
            'new_phone': new_phone
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
        user_info = UserInfo.objects.create(
            user=request.user,
            nickname=f'用户{request.user.username[:4]}',
            gender='未知',
            avatar=f'avatars/default_{avatar_num}.png',
            role='user'
        )
        # 确保传递角色信息
        return render(request, 'user/common/dashboard.html', {
            'user': request.user,
            'user_info': user_info,
            'user_info_with_role': {
                'role': getattr(request.user, 'role', 'user'),
                **user_info.__dict__
            }
        })

## 服务提供者仪表盘视图
## 功能: 显示服务提供者仪表盘, 包含服务信息和订单列表
def provider_dashboard(request):
    """服务提供者仪表盘视图"""
    if not request.user.is_authenticated:
        return redirect('login')
    if not hasattr(request.user, 'role') or request.user.role != 'provider':
        return redirect('user_dashboard')
    return render(request, 'provider\common\dashboard.html', {
        'user': request.user
    })

## 统一密码管理视图
## 功能: 支持忘记密码和修改密码
def unified_password_view(request, is_profile=False):
    """
    统一密码管理视图
    参数:
        is_profile: 是否来自个人资料页(True=修改密码, False=忘记密码)
    功能:
        1. 忘记密码: 通过手机验证码验证后重置
        2. 修改密码: 
           - 已设置密码: 验证原密码或手机验证码
           - 未设置密码: 直接设置新密码
    """
    User = get_user_model()
    
    if request.method == 'POST':
        # 获取表单数据
        phone = request.POST.get('phone')
        code = request.POST.get('verification_code', '')
        current_password = request.POST.get('current_password', '')
        new_password = request.POST.get('new_password')
        confirm_password = request.POST.get('confirm_password')
        
        # 验证密码一致性
        if not new_password or len(new_password) < 6 or len(new_password) > 12:
            messages.error(request, '密码长度需为6-12位')
            return redirect('password_reset')
        if not re.match(r'^(?=.*[A-Za-z])(?=.*\d)[A-Za-z\d]{6,12}$', new_password):
            messages.error(request, '密码需包含字母和数字组合')
            return redirect('password_reset')
            
        if new_password != confirm_password:
            messages.error(request, '两次输入的密码不一致')
            return redirect('password_reset')
        
        # 处理不同场景
        if is_profile and request.user.is_authenticated:
            # 个人资料页修改密码
            user = request.user
            
            # 检查用户是否有密码
            if user.has_usable_password():
                # 验证原密码或验证码
                if not (user.check_password(current_password) or 
                       (code and cache.get(f'reset_code_{user.phone}') == code)):
                    messages.error(request, '原密码错误或验证码无效')
                    return redirect('user_profile')
            # 未设置密码则直接更新
        else:
            # 忘记密码流程
            cached_code = cache.get(f'reset_code_{phone}')
            if not cached_code or cached_code != code:
                messages.error(request, '验证码错误或已过期')
                return redirect('password_reset')
                
            try:
                user = User.objects.get(phone=phone)
            except User.DoesNotExist:
                messages.error(request, '该手机号未注册')
                return redirect('password_reset')
        
        # 更新密码
        user.set_password(new_password)
        user.save()
        
        msg = '密码修改成功' if is_profile else '密码重置成功，请使用新密码登录'
        messages.success(request, msg)
        return redirect('user_profile' if is_profile else 'login')
    
    # GET请求处理
    context = {
        'is_profile': is_profile,
        'has_password': request.user.has_usable_password() if request.user.is_authenticated else False
    }
    template = 'user/change_password.html' if is_profile else 'registration/password_reset.html'
    return render(request, template, context)

## 用户个人资料视图
## 功能: 显示和编辑用户个人资料（头像、昵称、性别）
def profile_view(request):
    """用户个人资料视图"""
    if not request.user.is_authenticated:
        return redirect('login')
    
    try:
        user_info = UserInfo.objects.get(user=request.user)
    except UserInfo.DoesNotExist:
        # 如果UserInfo不存在，创建默认信息并设置用户角色
        avatar_num = random.randint(1, 6)
        user_info = UserInfo.objects.create(
            user=request.user,
            nickname=f'用户{request.user.username[:4]}',
            gender='未知',
            avatar=f'avatars/default_{avatar_num}.png',
            role='user'
        )
    
    if request.method == 'POST':
        # 处理表单提交
        try:
            # 更新昵称
            if 'nickname' in request.POST:
                nickname = request.POST['nickname'].strip()
                if 2 <= len(nickname) <= 10:
                    user_info.nickname = nickname
            
            # 更新性别
            if 'gender' in request.POST:
                gender = request.POST['gender']
                if gender in ['M', 'F', 'O']:
                    user_info.gender = gender
            
            # 处理头像上传
            if 'avatar' in request.FILES:
                avatar_file = request.FILES['avatar']
                # 验证文件类型和大小
                if avatar_file.content_type not in ['image/jpeg', 'image/png']:
                    raise ValueError('只支持JPEG/PNG格式图片')
                if avatar_file.size > 2*1024*1024:
                    raise ValueError('图片大小不能超过2MB')
                
                # 确保avatars目录存在
                avatars_dir = settings.MEDIA_ROOT
                os.makedirs(avatars_dir, exist_ok=True)
                
                # 生成唯一文件名
                ext = avatar_file.name.split('.')[-1].lower()
                filename = f'avatar_{request.user.id}_{int(time.time())}.{ext}'
                save_path = os.path.join(avatars_dir, filename)
                
                # 保存文件
                with open(save_path, 'wb+') as destination:
                    for chunk in avatar_file.chunks():
                        destination.write(chunk)
                
                # 更新头像路径
                user_info.avatar = filename
            
            user_info.save()
            
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': True,
                    'message': '个人信息更新成功',
                    'avatar_url': user_info.avatar.url if hasattr(user_info.avatar, 'url') else f'/media/{user_info.avatar}',
                    'nickname': user_info.nickname
                })
            messages.success(request, '个人信息更新成功')
        except Exception as e:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': False,
                    'message': str(e)
                }, status=400)
            messages.error(request, f'更新失败: {str(e)}')
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': False, 'message': '无效请求'}, status=400)
        return redirect('user_profile')
    
    return render(request, 'user/account/profile.html', {
        'user': request.user,
        'user_info': user_info
    })

#-----------------------------普通用户模块 安全设置-------------------------------#

## 安全设置视图
## 功能: 允许用户设置修改密码、绑定手机等安全选项
def security_settings(request):
    """安全设置视图"""
    if not request.user.is_authenticated:
        return redirect('login')
    try:
        user_info = UserInfo.objects.get(user=request.user)
    except UserInfo.DoesNotExist:
        # 如果UserInfo不存在，创建默认信息
        avatar_num = random.randint(1, 6)
        user_info = UserInfo.objects.create(
            user=request.user,
            nickname=f'用户{request.user.username[:4]}',
            gender='未知',
            avatar=f'avatars/default_{avatar_num}.png'
        )
    return render(request, 'user/account/security.html', {
        'user': request.user,
        'user_info': user_info
    })

## 设置密码API
## 功能: 处理设置密码请求, 支持JSON格式
def set_password_api(request):
    """设置密码API(用于未设置密码的用户)"""
    try:
        # 更精确的密码状态检查
        if request.user.password and request.user.has_usable_password():
            return JsonResponse({'status': 'error', 'message': '您已设置过密码，请使用修改密码功能'}, status=400)
        
        data = json.loads(request.body)
        new_password = data.get('new_password')
        confirm_password = data.get('confirm_password')
        
        # 验证必填字段
        if not new_password or not confirm_password:
            return JsonResponse({'status': 'error', 'message': '请填写所有密码字段'}, status=400)
            
        # 验证密码一致性
        if new_password != confirm_password:
            return JsonResponse({'status': 'error', 'message': '两次输入的密码不一致'}, status=400)
            
        # 验证密码复杂度
        if len(new_password) < 6 or len(new_password) > 12:
            return JsonResponse({'status': 'error', 'message': '密码长度需为6-12位'}, status=400)
        if not re.match(r'^(?=.*[A-Za-z])(?=.*\d)[A-Za-z\d]{6,12}$', new_password):
            return JsonResponse({'status': 'error', 'message': '密码需包含字母和数字组合'}, status=400)
            
        # 清除可能的无效密码哈希
        if request.user.password and not request.user.has_usable_password():
            request.user.password = ''
            
        # 设置新密码
        request.user.set_password(new_password)
        request.user.save()
        
        # 返回结果
        return JsonResponse({
            'status': 'success', 
            'message': '密码设置成功',
            'has_password': True,
            'redirect_url': '/user/security_settings/'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

## 修改密码API
## 功能: 处理修改密码请求, 支持JSON格式
@csrf_exempt
def change_password_api(request):
    """修改密码API"""
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': '仅支持POST请求'}, status=405)
    
    try:
        # 确保请求体是有效的JSON
        try:
            data = json.loads(request.body.decode('utf-8'))
        except json.JSONDecodeError:
            return JsonResponse({'status': 'error', 'message': '无效的JSON数据'}, status=400)
            
        old_password = data.get('old_password')
        new_password = data.get('new_password')
        confirm_password = data.get('confirm_password')
        
        # 验证必填字段
        if not all([old_password, new_password, confirm_password]):
            return JsonResponse({'status': 'error', 'message': '请填写所有密码字段'}, status=400)
            
        # 验证密码一致性
        if new_password != confirm_password:
            return JsonResponse({'status': 'error', 'message': '两次输入的新密码不一致'}, status=400)
            
        # 验证原密码
        if not request.user.check_password(old_password):
            return JsonResponse({'status': 'error', 'message': '原密码错误'}, status=400)
            
        # 验证密码复杂度
        if len(new_password) < 6 or len(new_password) > 12:
            return JsonResponse({'status': 'error', 'message': '密码长度需为6-12位'}, status=400)
        if not re.match(r'^(?=.*[A-Za-z])(?=.*\d)[A-Za-z\d]{6,12}$', new_password):
            return JsonResponse({'status': 'error', 'message': '密码需包含字母和数字组合'}, status=400)
            
        # 更新密码
        request.user.set_password(new_password)
        request.user.save()
        
        return JsonResponse({
            'status': 'success', 
            'message': '密码修改成功',
            'redirect_url': '/user/security_settings/'
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

def change_phone_view(request):
    """修改手机号视图"""
    if not request.user.is_authenticated:
        return redirect('login')
    
    if request.method == 'POST':
        new_phone = request.POST.get('new_phone')
        
        # 验证手机号格式
        if not re.match(r'^1[3-9]\d{9}$', new_phone):
            messages.error(request, '手机号格式不正确')
            return redirect('user_profile')
            
        # 更新手机号
        request.user.phone = new_phone
        request.user.save()
        messages.success(request, '手机号修改成功')
        return redirect('user_profile')
    
    return redirect('user_profile')

# 订单管理视图
def user_orders(request):
    """用户订单管理视图"""
    status = request.GET.get('status', 'active')
    return render(request, 'user/orders/orders.html', {
        'status': status,
        'orders': []  # 实际项目中这里应该查询数据库
    })

def quick_booking(request):
    """快速预约视图"""
    return render(request, 'user/quick_booking.html')

def favorites(request):
    """服务收藏夹视图"""
    return render(request, 'user/reviews/favorites.html')

def address_book(request):
    """地址簿视图"""
    return render(request, 'user/address/address_book.html')

def recommended_addresses(request):
    """推荐地址视图"""
    return render(request, 'user/recommended_addresses.html')

def refund_request(request):
    """退款申请视图"""
    return render(request, 'user/refund_request.html')

def rework_request(request):
    """返工申请视图"""
    return render(request, 'user/rework_request.html')

def complaint(request):
    """投诉建议视图"""
    return render(request, 'user/complaint.html')

def service_status(request):
    """服务进度查询视图"""
    return render(request, 'user/service_status.html')

def pending_reviews(request):
    """待评价订单视图"""
    return render(request, 'user/pending_reviews.html')

def review_history(request):
    """评价历史视图"""
    return render(request, 'user/review_history.html')

def featured_reviews(request):
    """精选评价视图"""
    return render(request, 'user/featured_reviews.html')



def provider_verification(request):
    """服务者认证视图"""
    return render(request, 'user/provider_verification.html')

#-----------------------------安全设置API-----------------------------#





def send_phone_code_api(request):
    """发送手机验证码API"""
    try:
        data = json.loads(request.body)
        phone = data.get('phone')
        
        if not phone or not re.match(r'^1[3-9]\d{9}$', phone):
            return JsonResponse({'status': 'error', 'message': '请输入有效的手机号码'}, status=400)
            
        # 生成6位随机验证码
        code = ''.join([str(random.randint(0, 9)) for _ in range(6)])
        # 存储验证码到缓存，有效期5分钟
        cache.set(f'security_code_{phone}', code, 300)
        
        return JsonResponse({
            'status': 'success',
            'message': '验证码已发送',
            'code': code  # 开发模式下返回验证码
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

def change_phone_api(request):
    """更换手机号API(需验证原手机号验证码)"""
    try:
        data = json.loads(request.body)
        old_phone = request.user.phone  # 获取当前绑定手机号
        new_phone = data.get('new_phone')
        code = data.get('code')
        
        # 验证新手机号格式
        if not new_phone or not re.match(r'^1[3-9]\d{9}$', new_phone):
            return JsonResponse({'status': 'error', 'message': '请输入有效的新手机号码'}, status=400)
            
        # 验证码检查
        if not code or not re.match(r'^\d{6}$', code):
            return JsonResponse({'status': 'error', 'message': '请输入6位验证码'}, status=400)
            
        # 验证原手机号验证码
        cached_code = cache.get(f'change_phone_code_{old_phone}')
        if not cached_code or cached_code != code:
            return JsonResponse({'status': 'error', 'message': '验证码错误或已过期'}, status=400)
            
        # 检查新手机号是否已被使用
        User = get_user_model()
        if User.objects.filter(phone=new_phone).exclude(id=request.user.id).exists():
            return JsonResponse({'status': 'error', 'message': '该手机号已被其他账号使用'}, status=400)
            
        # 更新手机号
        request.user.phone = new_phone
        request.user.save()
        
        # 清除验证码缓存
        cache.delete(f'change_phone_code_{old_phone}')
        
        return JsonResponse({
            'status': 'success', 
            'message': '手机号更换成功',
            'new_phone': new_phone  # 返回新手机号供前端更新显示
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@csrf_exempt
def send_change_phone_code(request):
    """发送更换手机号验证码API(开发模式)"""
    if request.method != 'POST':
        return JsonResponse({'status': 'error', 'message': '仅支持POST请求'}, status=405)
        
    try:
        # 获取当前用户手机号
        if not request.user.is_authenticated:
            return JsonResponse({'status': 'error', 'message': '用户未登录'}, status=401)
            
        phone = request.user.phone
        if not phone:
            return JsonResponse({'status': 'error', 'message': '未绑定手机号'}, status=400)
            
        # 生成6位随机验证码
        code = ''.join([str(random.randint(0, 9)) for _ in range(6)])
        # 存储验证码到缓存，有效期5分钟
        cache.set(f'change_phone_code_{phone}', code, 300)
        
        return JsonResponse({
            'status': 'success',
            'message': '验证码已发送(开发模式)',
            'code': code,  # 开发模式下返回验证码
            'phone': phone  # 返回当前绑定手机号
        })
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

def submit_auth_api(request):
    """实名认证API"""
    try:
        data = json.loads(request.body)
        real_name = data.get('real_name')
        id_number = data.get('id_number')
        
        if not real_name or len(real_name) < 2:
            return JsonResponse({'status': 'error', 'message': '请输入有效的真实姓名'}, status=400)
            
        if not id_number or not re.match(r'(^\d{15}$)|(^\d{17}(\d|X|x)$)', id_number):
            return JsonResponse({'status': 'error', 'message': '请输入有效的身份证号码'}, status=400)
            
        # 更新用户实名信息
        request.user.real_name = real_name
        request.user.id_number = id_number
        request.user.save()
        return JsonResponse({'status': 'success', 'message': '实名认证成功'})
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

def delete_account_api(request):
    """账号注销API"""
    try:
        # 标记用户为已删除
        request.user.is_active = False
        request.user.save()
        
        # 登出用户
        from django.contrib.auth import logout
        logout(request)
        
        return JsonResponse({'status': 'success', 'message': '账号已注销'})
        
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

def password_status_api(request):
    """密码状态检查API"""
    try:
        if not request.user.is_authenticated:
            return JsonResponse(
                {'status': 'error', 'message': '用户未登录'}, 
                status=401
            )
            
        return JsonResponse({
            'status': 'success',
            'has_password': request.user.has_usable_password(),
            'message': '密码状态获取成功'
        })
        
    except Exception as e:
        return JsonResponse(
            {'status': 'error', 'message': str(e)},
            status=500
        )
